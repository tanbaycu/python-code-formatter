import autopep8
import os
import logging
import ast
import shutil
import subprocess
from rich.console import Console
from rich.syntax import Syntax
from rich.panel import Panel
from rich.prompt import Prompt
from rich.text import Text
from rich import box
from docx import Document
import pyperclip
import time
import tempfile
import pygments
import pygments.formatters.html
import pygments.lexers
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from webdriver_manager.chrome import ChromeDriverManager
from radon.complexity import cc_visit
from radon.visitors import ComplexityVisitor


# Setup logging
logging.basicConfig(filename='code_formatter.log', level=logging.ERROR,
                    format='%(asctime)s - %(levelname)s - %(message)s')

console = Console()


def format_code(code):
    try:
        formatted_code = autopep8.fix_code(code, options={'aggressive': 2})
        return formatted_code
    except Exception as e:
        console.print(f"[bold red]Error: {e}[/bold red]")
        logging.error(f"Error formatting code: {e}")
        return None


def save_to_file(code, filename):
    try:
        with open(filename, 'w', encoding='utf-8') as file:
            file.write(code)
        console.print(
            f"[bold green]:white_check_mark: Code has been saved to file {filename}[/bold green]")
        console.print(
            f"[bold yellow]File path: {os.path.abspath(filename)}[/bold yellow]")
    except Exception as e:
        console.print(f'[bold red]Unable to save code: {e}[/bold red]')
        logging.error(f'Unable to save code to file {filename}: {e}')


def copy_to_clipboard(code):
    try:
        pyperclip.copy(code)
        console.print(
            "[bold green]:white_check_mark: Code has been copied to clipboard.[/bold green]")
    except Exception as e:
        console.print(f'[bold red]Unable to copy code: {e}[/bold red]')
        logging.error(f'Unable to copy code to clipboard: {e}')


def export_to_image(code, filename):
    try:
        lexer = pygments.lexers.get_lexer_by_name("python")
        formatter = pygments.formatters.html.HtmlFormatter(style="dracula")
        html_code = pygments.highlight(code, lexer, formatter)

        css = formatter.get_style_defs('.highlight')
        html_code = f"""
        <html>
        <head>
            <style>
                body {{
                    background-color: #1e1e2e;
                    color: #e0e0e0;
                    margin: 0;
                    padding: 0;
                    font-family: 'Source Code Pro', monospace;
                }}
                .highlight {{
                    background-color: #1e1e2e;
                    border-radius: 8px;
                    padding: 20px;
                    overflow: auto;
                }}
                pre {{
                    margin: 0;
                }}
                {css}
            </style>
        </head>
        <body>
            <div class="highlight">
                <pre>{html_code}</pre>
            </div>
        </body>
        </html>
        """

        with tempfile.NamedTemporaryFile(delete=False, suffix=".html") as temp_html:
            temp_html.write(html_code.encode('utf-8'))
            temp_html_path = temp_html.name

        chrome_options = webdriver.ChromeOptions()
        chrome_options.add_argument("--headless")
        chrome_options.add_argument("--disable-gpu")

        with webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options) as driver:
            driver.get(f"file://{temp_html_path}")

            content = driver.find_element(By.TAG_NAME, 'pre')
            content_height = content.size['height']
            content_width = content.size['width']

            # Adjust size if code is too long
            driver.set_window_size(content_width + 500, content_height + 500)

            driver.save_screenshot(filename)

        os.remove(temp_html_path)

        console.print(
            f"[bold green]:white_check_mark: Code has been exported to image file {filename}[/bold green]")

    except Exception as e:
        console.print(f'[bold red]Unable to export image: {e}[/bold red]')
        logging.error(f'Unable to export image {filename}: {e}')


def export_to_word(code, filename):
    try:
        if not filename.lower().endswith('.docx'):
            filename += '.docx'
        doc = Document()
        doc.add_heading('Code', level=1)
        doc.add_paragraph(code)
        doc.save(filename)
        console.print(
            f"[bold green]:white_check_mark: Code has been exported to Word file {filename}[/bold green]")
    except Exception as e:
        console.print(f'[bold red]Unable to export to Word: {e}[/bold red]')
        logging.error(f'Unable to export to Word file {filename}: {e}')


def export_to_document(code, filename):
    try:
        if not filename.lower().endswith('.md'):
            filename += '.md'
        with open(filename, 'w', encoding='utf-8') as file:
            file.write(f'```python\n{code}\n```')
        console.print(
            f"[bold green]:white_check_mark: Code has been exported to Markdown document {filename}[/bold green]")
    except Exception as e:
        console.print(f'[bold red]Unable to export document: {e}[/bold red]')
        logging.error(f'Unable to export Markdown document {filename}: {e}')


def count_lines(code):
    """Count the number of lines in the code."""
    return len(code.splitlines())


def count_characters(code):
    """Count the number of characters in the code."""
    return len(code)


def time_formatting(func, *args):
    """Measure the time taken to format code."""
    start_time = time.time()
    result = func(*args)
    end_time = time.time()
    duration = end_time - start_time
    return result, duration


def analyze_dependencies(code):
    """Analyze code to identify imported libraries."""
    tree = ast.parse(code)
    dependencies = set()

    for node in ast.walk(tree):
        if isinstance(node, ast.Import):
            for alias in node.names:
                dependencies.add(alias.name)
        elif isinstance(node, ast.ImportFrom):
            dependencies.add(node.module)

    return dependencies


def show_dependencies(code):
    """Display the dependencies of the code."""
    dependencies = analyze_dependencies(code)
    console.print("[bold cyan]Source code dependencies:[/bold cyan]")
    if dependencies:
        console.print(
            f"[bold green]Imported libraries:[/bold green] {', '.join(dependencies)}")
    else:
        console.print("[bold red]No dependencies found.[/bold red]")


def calculate_code_complexity(code):
    """Calculate the cyclomatic complexity of the code."""
    try:
        tree = ast.parse(code)
        complexities = cc_visit(tree)
        max_complexity = max(c.complexity for c in complexities)
        return max_complexity
    except Exception as e:
        console.print(
            f"[bold red]Error calculating complexity: {e}[/bold red]")
        return None


def find_unused_code_segments(code):
    """Find unused code segments like variables and functions."""
    tree = ast.parse(code)
    used_names = set()

    class NameVisitor(ast.NodeVisitor):
        def visit_Name(self, node):
            used_names.add(node.id)

    class FunctionVisitor(ast.NodeVisitor):
        def visit_FunctionDef(self, node):
            used_names.add(node.name)

    class ClassVisitor(ast.NodeVisitor):
        def visit_ClassDef(self, node):
            used_names.add(node.name)

    NameVisitor().visit(tree)
    FunctionVisitor().visit(tree)
    ClassVisitor().visit(tree)

    all_names = {node.id for node in ast.walk(
        tree) if isinstance(node, ast.Name)}
    all_functions = {node.name for node in ast.walk(
        tree) if isinstance(node, ast.FunctionDef)}
    all_classes = {node.name for node in ast.walk(
        tree) if isinstance(node, ast.ClassDef)}

    unused_names = all_names - used_names
    unused_functions = all_functions - used_names
    unused_classes = all_classes - used_names

    return {
        "unused_variables": list(unused_names),
        "unused_functions": list(unused_functions),
        "unused_classes": list(unused_classes)
    }


def display_unused_code_segments(code):
    """Display unused code segments."""
    unused_code = find_unused_code_segments(code)

    console.print(f"[bold cyan]Unused code segments:[/bold cyan]")
    if unused_code["unused_variables"]:
        console.print(
            f"[bold red]Unused variables:[/bold red] {', '.join(unused_code['unused_variables'])}")
    else:
        console.print("[bold green]No unused variables detected.[/bold green]")

    if unused_code["unused_functions"]:
        console.print(
            f"[bold red]Unused functions:[/bold red] {', '.join(unused_code['unused_functions'])}")
    else:
        console.print("[bold green]No unused functions detected.[/bold green]")

    if unused_code["unused_classes"]:
        console.print(
            f"[bold red]Unused classes:[/bold red] {', '.join(unused_code['unused_classes'])}")
    else:
        console.print("[bold green]No unused classes detected.[/bold green]")


def generate_detailed_report(code):
    """Generate a detailed report about the code."""
    tree = ast.parse(code)
    report = {
        "classes": 0,
        "functions": 0,
        "variables": 0,
    }

    class DetailedReportVisitor(ast.NodeVisitor):
        def visit_ClassDef(self, node):
            report["classes"] += 1
            self.generic_visit(node)

        def visit_FunctionDef(self, node):
            report["functions"] += 1
            self.generic_visit(node)

        def visit_Name(self, node):
            if isinstance(node.ctx, ast.Store):
                report["variables"] += 1
            self.generic_visit(node)

    DetailedReportVisitor().visit(tree)

    return report


def display_detailed_report(code):
    """Display a detailed report about the code."""
    report = generate_detailed_report(code)

    console.print(f"[bold cyan]Detailed Code Report:[/bold cyan]")
    console.print(
        f"[bold #000080]Number of classes:[/bold #000080] [bold D2B48C]{report['classes']}[/bold D2B48C]")
    console.print(
        f"[bold #000080]Number of functions:[/bold #000080] [bold D2B48C]{report['functions']}[/bold D2B48C]")
    console.print(
        f"[bold #000080]Number of variables:[/bold #000080] [bold D2B48C]{report['variables']}[/bold D2B48C]")


if __name__ == "__main__":
    while True:
        console.print(
            ":pencil: [bold cyan]Enter your code (press Ctrl+D or Ctrl+Z => Enter to end):[/bold cyan]")

        user_code = ""
        try:
            while True:
                line = input()
                user_code += line + "\n"
        except EOFError:
            pass

        line_count_before = count_lines(user_code)
        char_count = count_characters(user_code)

        formatted_code = format_code(user_code)

        if formatted_code:
            console.clear()

            line_count_after = count_lines(formatted_code)

            original_panel = Panel(
                Syntax(user_code, "python", theme="monokai", line_numbers=True),
                title="Original Code",
                title_align="left",
                border_style="bold cyan",
                box=box.ROUNDED
            )
            formatted_panel = Panel(
                Syntax(formatted_code, "python",
                       theme="monokai", line_numbers=True),
                title="Formatted Code",
                title_align="left",
                border_style="bold green",
                box=box.ROUNDED,
                subtitle=":clipboard: Press 'c' to copy code",
                subtitle_align="right"
            )

            console.print(original_panel)
            console.print(formatted_panel)

            console.print(
                f"[bold cyan]Number of lines before: {line_count_before}[/bold cyan]")
            console.print(
                f"[bold cyan]Number of lines after: {line_count_after}[/bold cyan]")
            console.print(
                f"[bold cyan]Number of characters in code: {char_count}[/bold cyan]")

            formatted_code, duration = time_formatting(
                format_code, user_code)
            console.print(
                f"[bold cyan]Time taken to format code: {duration:.2f} seconds[/bold cyan]")

            console.print(
                f"[bold #FF1493]Lines of code before: [bold #9ACD32]{line_count_before}[/bold #FF1493]"
                f"\n[bold #FF1493]Lines of code after: [bold #9ACD32]{line_count_after}[/bold #FF1493]"
                f"\n[bold #FF1493]Characters in code: [bold #9ACD32]{char_count}[/bold #FF1493]"
            )

            show_dependencies(user_code)

            complexity = calculate_code_complexity(user_code)
            if complexity is not None:
                console.print(
                    f"Cyclomatic complexity of the code: {complexity}")

            display_unused_code_segments(user_code)

            display_detailed_report(user_code)

            while True:
                user_input = Prompt.ask(
                    "[bold yellow]Press:\n"
                    "'c' to copy the formatted code\n"
                    "'s' to save the code to a file\n"
                    "'p' to export the code to an image or Word\n"
                    "Or any other key to skip[/bold yellow]").strip().lower()

                if user_input == 'c':
                    copy_to_clipboard(formatted_code)
                elif user_input == 's':
                    filename = Prompt.ask(
                        ":floppy_disk: [bold yellow]Enter the filename to save the code:[/bold yellow]").strip()
                    save_to_file(formatted_code, filename)
                elif user_input == 'p':
                    export_choice = Prompt.ask(
                        "[bold yellow]Press:\n"
                        "'i' to export the code to an image\n"
                        "'w' to export the code to a Word file[/bold yellow]").strip().lower()

                    if export_choice == 'i':
                        image_filename = Prompt.ask(
                            ":floppy_disk: [bold yellow]Enter the image filename (e.g., code.png):[/bold yellow]").strip()
                        export_to_image(
                            formatted_code, image_filename)
                    elif export_choice == 'w':
                        word_filename = Prompt.ask(
                            ":floppy_disk: [bold yellow]Enter the Word filename (e.g., code.docx):[/bold yellow]").strip()
                        export_to_word(formatted_code, word_filename)
                    else:
                        console.print(
                            "[bold red]Invalid choice.[/bold red]")
                        logging.error(
                            f'Invalid file format choice: {export_choice}')
                else:
                    break

        user_input = Prompt.ask(
            ":question: [bold cyan]Do you want to continue with another code format? (y/n)[/bold cyan]").strip().lower()
        if user_input == 'n':
            console.print("[bold red]:wave: Goodbye[/bold red]")
            break
        elif user_input == 'y':
            console.clear()
        else:
            console.print(
                "[bold red]:x: Invalid selection. Please select 'y' or 'n'.[/bold red]")
            logging.error(f'Invalid selection: {user_input}')

